import os
from typing import Text
from urllib import request
import docx
from simplify_docx import simplify
from pathlib import Path
import re
from win32com import client as wc
import requests
import mechanize

doc = docx.Document('D:/programming/python/automate/data folder/घर बिवरण 200.docx')

house_number_regx = re.compile('\d+')
def get_house_number():
    house_number = Path('D:/programming/python/automate/घर बिवरण 200.docx').name
    
    n = house_number_regx.search(house_number)
    return n.group()

house_number = get_house_number()
print(house_number)



text = []
for p in doc.paragraphs:
    text.append(p.text)




doc_text = text[1].replace("\t","")
doc_text2 = doc_text.replace(" ","")
doc_text = doc_text2



addressRegex = re.compile('(?<=ठेगाना:).*,(\s)?\d+')
address = addressRegex.search(doc_text)
print('Address:')
try:
    print(address.group())
except AttributeError:
    address = '-'
    print(address)



toleRegex = re.compile('टोल(.*)धर्म')
tole = toleRegex.search(doc_text)
print('Tole:')
try:
    tole = tole.groups()[0].strip(":")
except AttributeError:
    tole = "-"
    





religionRegex = re.compile('धर्म(.*)मो')
religion = religionRegex.search(doc_text)
print('Religion:')
try:
    religion = religion.groups()[0].strip(":")
except AttributeError:
    religion = "-"
    






mobileRegex = re.compile(r'\d{10}')
mobile = mobileRegex.search(doc_text)
print('Mobile:')
try:
    mobile = mobile.group()
except AttributeError:
    mobile = '-'
    



print("Members:")
family_persons = []
i = 1

count_row = 0

for row in doc.tables[0].rows:
    count_row +=1

total_people = count_row - 1

while i <=7:
    name = doc.tables[0].rows[1].cells[i].text
    family_persons.append(name)
    i +=1


a = 2
while a <= total_people:
    c = 1
    while c <= 7:
        name = doc.tables[0].rows[a].cells[c].text
        if name == "Æ":
            name = "माथिकै"
        elif name == '"':
            name ="माथिकै"
        elif name == "":
            name = "-"
        else:
            name = doc.tables[0].rows[a].cells[c].text
        family_persons.append(name)
        c += 1
    a += 1 








economic_status = []
cell_counter = 0
while cell_counter <=4:
    values = doc.tables[1].rows[1].cells[cell_counter].text
    if values == "":
        values = "-"
    economic_status.append(values)
    cell_counter +=1
print('Economic:')
print(economic_status)




living_status =[]

land_status = ""
land_area =""
if doc.tables[2].rows[1].cells[0].text == "-":
    land_status = "सुकुम्बासी वा बहालमा"
elif doc.tables[2].rows[1].cells[0].text == "":
    land_status = "सुकुम्बासी वा बहालमा"
elif doc.tables[2].rows[1].cells[0].text == "छैन":
    land_status = "सुकुम्बासी वा बहालमा"    
else:
    land_status = "भएकाे"
    land_area = doc.tables[2].rows[1].cells[0].text
print('Land Status:')
print(land_status)
print('Land Area:')
print(land_area)

house_type =""


doc_as_json = simplify(doc)


print('House type:')
try:
    var = doc_as_json['VALUE'][0]['VALUE'][6]['VALUE'][2]['VALUE'][1]['VALUE'][0]['VALUE'][0]['VALUE']
    house_type = "पक्की"
except IndexError:
    try:
        var = doc_as_json['VALUE'][0]['VALUE'][6]['VALUE'][2]['VALUE'][2]['VALUE'][0]['VALUE'][0]['VALUE']
        house_type = "अर्रधपक्की"
    except IndexError:
        try:
            var = doc_as_json['VALUE'][0]['VALUE'][6]['VALUE'][2]['VALUE'][3]['VALUE'][0]['VALUE'][0]['VALUE']
            house_type = "काठ"
        except IndexError:
            try:
                var = doc_as_json['VALUE'][0]['VALUE'][6]['VALUE'][2]['VALUE'][4]['VALUE'][0]['VALUE'][0]['VALUE']
                house_type = "कच्ची"
            except IndexError:
                house_type = "अज्ञात"

print(house_type)

number_of_rooms = ""
print("Number of Rooms:")

print(number_of_rooms)




vehicle_status = doc.tables[3].rows[1].cells[0].text

print('Vechicle Status:')
print(vehicle_status)



comm_status = str(doc.tables[3].rows[1].cells[1].text)
print('Communication Status:')



print(comm_status)





refrigerator = doc.tables[3].rows[1].cells[2].text
print("Refrigerator:")
print(refrigerator)

animals = doc.tables[3].rows[1].cells[3].text

print("Domestic animals:")
print(animals)

birds = doc.tables[3].rows[1].cells[4].text

print("Domestic birds:")
print(birds)

water_source = doc.tables[4].rows[1].cells[0].text
print("Water Source:")
print(water_source)


toilet_facility_type = doc.tables[4].rows[1].cells[1].text
cooking_fuel = doc.tables[4].rows[1].cells[2].text
light_source = doc.tables[4].rows[1].cells[3].text

print("Toilet Facility & Type:")
print(toilet_facility_type)

print("Cooking Fuel:")
print(cooking_fuel)

print("Light Source:")
print(light_source)

geographical_status = doc.tables[5].rows[1].cells[0].text
health_service = doc.tables[5].rows[1].cells[1].text
land_value = doc.tables[5].rows[1].cells[2].text
road_facility = doc.tables[5].rows[1].cells[3].text
remarks = doc.tables[5].rows[1].cells[4].text

print("Geographical status:")
print(geographical_status)

print("Health Service:")
print(health_service)

print("Land Value:")
print(land_value)

print("Road Facility:")
print(road_facility)

print("Remarks: ")
print(remarks)


head = family_persons[0]
head_gender = family_persons[1]
head_age = family_persons[2]
head_relation = family_persons[3]
head_citizenship = family_persons[4]
head_occupation = family_persons[5]
head_education = family_persons[6]


print(family_persons)


total_fam_arr_data = len(family_persons)
all_mem = family_persons[7:total_fam_arr_data]
print(all_mem)